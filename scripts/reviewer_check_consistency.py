import docx, re, os
os.chdir(r"C:/Users/12751/WorkBuddy/脓毒症多器官轨迹")
doc = docx.Document('脓毒症多器官轨迹_manuscript_BMC.docx')
docx_text = '\n'.join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            docx_text += '\n' + c.text
with open('P7_manuscript_clean.md', encoding='utf-8') as f:
    md_text = f.read()
md_clean = re.sub(r'[#*|`>]', '', md_text)
md_clean = re.sub(r'\n{2,}', '\n', md_clean)
tokens = ['24,098','34,003','24,101','3,680','15.3','27.9','24.5','16.4','15.9','15.3','7.7','37.6',
'0.56','0.50','0.70','1.44','1.25','0.51','0.39','1.62','0.65','0.73','0.76',
'15,292','0.42','0.81','5,358','12,055','1,037','4.3','67.8','58.2',
'0.827','0.838','0.971','0.446','0.890','13.05','0.869','0.161','0.125',
'4.87','2.61','0.004','0.014','QNPY2023-30','1275128366','0000-0003-0875-0837','15390234035']
missing_in_docx = [t for t in tokens if t not in docx_text]
missing_in_md = [t for t in tokens if t not in md_clean]
print('Missing in DOCX:', missing_in_docx if missing_in_docx else 'NONE - all present')
print('Missing in MD:', missing_in_md if missing_in_md else 'NONE - all present')
