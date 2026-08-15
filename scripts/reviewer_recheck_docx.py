import docx, re
doc = docx.Document('脓毒症多器官轨迹_manuscript_BMC.docx')
full = '\n'.join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            full += '\n' + c.text

print("=== DOCX placeholder scan ===")
for pat in ['To be confirmed', 'placeholder', 'PLACEHOLDER', '(planned)', 'planned)', 'Introduction', '~4,900', 'Figure to be generated', 'XXX', 'TBD']:
    n = len(re.findall(re.escape(pat), full, re.I))
    print(f"  {pat}: {n}")

print("\n=== DOCX new elements ===")
print("  Background section:", '## Background' in full or 'Background' in full[:6000])
print("  ~4,500 wordcount:", '~4,500' in full)
print("  Additional file 1:", 'Additional file 1' in full)
print("  Additional file 2:", 'Additional file 2' in full)
print("  Additional file 3:", 'Additional file 3' in full)
print("  ethics exemption wording:", 'determined to be exempt' in full)
print("  GitHub/deposit wording:", 'public repository' in full or 'GitHub' in full)

# word count check
print("\n=== Word count sanity (from DOCX paragraphs Introduction..Conclusions via Background..Conclusions) ===")
paras = [p.text for p in doc.paragraphs]
# count from Background heading to List of Abbreviations
start = None; end = None
for i, t in enumerate(paras):
    if t.strip() == 'Background' and start is None:
        start = i
    if t.strip() == 'List of Abbreviations':
        end = i
        break
if start is not None and end is not None:
    seg = ' '.join(paras[start+1:end])
    print(f"  Background->Conclusions words (incl headings in seg): {len(seg.split())}")
else:
    print("  could not locate boundaries:", start, end)
