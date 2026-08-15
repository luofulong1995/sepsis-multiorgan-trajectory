import docx
from docx.oxml.ns import qn
doc = docx.Document('脓毒症多器官轨迹_manuscript_BMC.docx')

# 1. Normal style font & spacing
st = doc.styles['Normal']
print("Normal font:", st.font.name, st.font.size)
print("Normal line spacing rule:", st.paragraph_format.line_spacing_rule)

# 2. Line numbers in sectPr
sectPr = doc.sections[0]._sectPr
ln = sectPr.find(qn('w:lnNumType'))
print("Line numbering present:", ln is not None, "| restart:", ln.get(qn('w:restart')) if ln is not None else None)

# 3. Page number in footer
footer_paras = doc.sections[0].footer.paragraphs
xml = footer_paras[0]._p.xml if footer_paras else ''
print("Footer has PAGE field:", 'PAGE' in xml)

# 4. Margins
sec = doc.sections[0]
print("Margins (in): top", sec.top_margin.inches, "bottom", sec.bottom_margin.inches, "left", sec.left_margin.inches, "right", sec.right_margin.inches)

# 5. Check headings used (should not be too many level issues)
print("Sample heading style names:", sorted(set(p.style.name for p in doc.paragraphs if p.style.name.startswith('Heading'))))

# 6. Check table title above each table
print("\nTable count:", len(doc.tables))
